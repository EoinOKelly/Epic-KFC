"""Build CS4455 cryptographic design document as Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "Cryptographic-Design-kfc.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = text
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    set_normal_style(doc)

    title = doc.add_heading("Cryptographic Design Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "CS4455 Cybersecurity Epic Project 2026\n"
        "Team kfc — Epic Messaging\n"
        "Module: Cryptography (Eoin O'Brien)"
    )
    run.font.size = Pt(12)

    doc.add_paragraph(
        "Epic Messaging is a secure relay messaging system. Message bodies are end-to-end encrypted "
        "before they reach the server. Passwords and long-term private keys are protected with "
        "standard KDFs and AEAD. The Qt client uses OpenSSL for production traffic; the cryptography/ "
        "package implements the full Signal-style session for automated tests and backend integration."
    )

    doc.add_heading("1. Security goals", level=1)
    doc.add_paragraph(
        "The design targets confidentiality, integrity, and authenticity of message content between users."
    )
    add_table(
        doc,
        ["Property", "Mechanism"],
        [
            ["Confidentiality of message bodies", "AES-256-GCM inside an E2EE session; server stores ciphertext only"],
            ["Integrity of message bodies", "GCM authentication tag; tampering fails decryption"],
            ["Authenticity of sender", "Ed25519-signed pre-keys plus TOFU identity pinning on clients"],
            ["Password confidentiality if DB leaks", "Argon2id PHC hashes; no plaintext passwords"],
            ["Private keys at rest on client", "HKDF-derived keys and AES-256-GCM via encryptPrivateKeyForStorage"],
        ],
    )
    doc.add_paragraph(
        "TLS protects the transport to the API host. E2EE protects content from the server operator."
    )

    doc.add_heading("2. Threat model", level=1)
    doc.add_paragraph(
        "We consider four attacker classes. The tables below separate cryptographic guarantees "
        "from properties the protocol does not provide."
    )
    add_table(
        doc,
        ["Attacker", "Capability"],
        [
            ["A. Passive network", "Read client-server traffic"],
            ["B. Active network", "Modify, drop, replay, or inject traffic"],
            ["C. Honest-but-curious server", "Follows the API; logs ciphertext and metadata"],
            ["D. Compromised server", "Full database; can return malicious API responses"],
        ],
    )

    doc.add_heading("2.1 Cryptographic defences", level=2)
    add_table(
        doc,
        ["Defence", "A", "B", "C", "D"],
        [
            ["Past ciphertext stays confidential", "Yes", "Yes", "Yes", "Yes"],
            ["Ciphertext tampering is detected", "n/a", "Yes", "Yes", "Yes"],
            ["Sender cryptographically bound to keys", "n/a", "Yes", "Yes", "Yes*"],
            ["Forward secrecy for new messages", "Yes", "Yes", "Yes", "Yes**"],
        ],
    )
    doc.add_paragraph(
        "*Clients pin identity keys with /trust and block sends if the server swaps keys."
    )
    doc.add_paragraph(
        "**Full double ratchet in the TypeScript package; the C++ client runs X3DH plus GCM on each message "
        "until persisted ratchet state is added."
    )

    doc.add_heading("2.2 Out of scope", level=2)
    doc.add_paragraph(
        "These properties are not provided by the cipher layer or are handled operationally:"
    )
    add_table(
        doc,
        ["Topic", "Status", "Explanation"],
        [
            ["Hide who messaged whom", "Not provided", "Relay server and TLS observer see routing metadata"],
            ["Guarantee message delivery", "Not provided", "Server can drop or withhold messages"],
            ["First-contact MITM", "Mitigated by TOFU", "User runs /trust before sending"],
        ],
    )

    doc.add_heading("2.3 Compromised server", level=2)
    doc.add_paragraph("With full database access, the server still cannot:", style="List Bullet")
    doc.add_paragraph("Decrypt wire_payload_json without client private keys.", style="List Bullet")
    doc.add_paragraph("Forge AEAD ciphertext that decrypts to chosen plaintext for an existing session.", style="List Bullet")
    doc.add_paragraph(
        "The server can observe metadata, delay delivery, or substitute pre-keys for contacts not yet pinned.",
        style="List Bullet",
    )

    doc.add_heading("3. Protocol construction", level=1)

    doc.add_heading("3.1 Registration and key publication", level=2)
    doc.add_paragraph(
        "1. Client generates X25519 and Ed25519 key pairs and pre-keys.\n"
        "2. Client registers over TLS; server stores Argon2id password hash only.\n"
        "3. Client uploads public pre-key bundle; private keys remain on device.\n"
        "4. Private key blobs may be encrypted at rest under HKDF-derived keys."
    )

    doc.add_heading("3.2 Send and receive", level=2)
    doc.add_paragraph(
        "1. Sender fetches recipient public bundle after /trust pins identity.\n"
        "2. TypeScript (cryptography/): X3DH, double ratchet, GCM envelope, libsignal-v1 wire JSON.\n"
        "3. C++ client: OpenSSL X3DH and GCM per message; each send carries a full X3DH block today.\n"
        "4. Server stores wire_payload_json; recipient decrypts locally.\n"
        "5. Backend checks JSON shape only; it never decrypts payloads."
    )

    doc.add_heading("3.3 Trust model", level=2)
    doc.add_paragraph(
        "Public keys are stored on the server. Clients pin the identity key on first use (/trust) "
        "and stop sending if the identity changes. We use TOFU rather than a certificate authority."
    )

    doc.add_heading("4. Primitive choices", level=1)
    doc.add_paragraph(
        "Libraries: Node crypto, argon2, @privacyresearch/libsignal-protocol-typescript; "
        "C++ client: OpenSSL 3 EVP."
    )

    doc.add_heading("4.1 Argon2id", level=2)
    add_table(
        doc,
        ["Parameter", "Value", "Reference"],
        [
            ["Type", "Argon2id", "RFC 9106"],
            ["Memory", "64 MiB", "OWASP Password Storage Cheat Sheet 2024 (strong tier)"],
            ["Iterations", "3", "OWASP 2024"],
            ["Parallelism", "4", "OWASP 2024"],
        ],
    )

    doc.add_heading("4.2 HKDF-SHA256 and AES-256-GCM", level=2)
    doc.add_paragraph(
        "HKDF (RFC 5869) uses separate info strings for local storage keys versus session keys. "
        "AES-256-GCM (NIST SP 800-38D) uses a random 12-byte nonce per encryption. Reusing a nonce "
        "under the same key breaks confidentiality; we use random IVs and ratchet-derived message keys."
    )

    doc.add_heading("4.3 X3DH and session keys", level=2)
    doc.add_paragraph(
        "HPKE Mode_Auth (RFC 9180) is one standard option for authenticated key establishment. "
        "We use X3DH (Signal, 2016) with pre-key bundles on the relay server. The TypeScript package "
        "adds the double ratchet (Signal, 2016) for ongoing messages. X25519 (RFC 7748) and Ed25519 "
        "(RFC 8032) handle agreement and signed pre-keys."
    )

    doc.add_heading("5. Implementation", level=1)
    add_table(
        doc,
        ["Component", "Role"],
        [
            ["cryptography/src/cryptoEngine.ts", "Argon2id, HKDF, AES-GCM helpers"],
            ["cryptography/src/signal/", "X3DH, double ratchet, TOFU, libsignal-v1 wire"],
            ["client NativeSignalCryptoProvider", "OpenSSL: keys, pre-key verify, X3DH+GCM per message"],
            ["FastAPI backend", "Relay; Argon2id passwords; no payload decryption"],
        ],
    )

    doc.add_heading("6. Limitations", level=1)
    for item in [
        "Metadata (participants, timing) is visible to the server.",
        "The C++ client does not yet persist double-ratchet state (see client README).",
        "C++ and TypeScript inner wire formats are not interoperable without a shared adapter.",
        "X25519 is not post-quantum; Kyber/PQXDH is reserved for a future libsignal upgrade.",
        "Users should run /trust before the first message to a new contact.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Validation", level=1)
    add_table(
        doc,
        ["Check", "Command"],
        [
            ["E2EE round-trip (TypeScript)", "cd cryptography && npm run smoke:signal"],
            ["API relay", "npm run e2e:backend"],
            ["C++ X3DH + GCM", "client_tests with OpenSSL enabled"],
        ],
    )

    doc.add_heading("References", level=1)
    for ref in [
        "RFC 5869, RFC 7748, RFC 8032, RFC 9106, RFC 9180",
        "NIST SP 800-38D, NIST SP 800-57",
        "OWASP Password Storage Cheat Sheet (2024)",
        "Signal X3DH and Double Ratchet specifications (2016)",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
